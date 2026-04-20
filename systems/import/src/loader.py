"""Document loader for the RAG Data Import pipeline.

Loads raw documents from files (PDF, TXT, MD, DOCX), HTTP/HTTPS URLs, and
directories. All loaders return a list of RawDocument objects.
"""

import logging
from pathlib import Path

from src.models import RawDocument

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}


class DocumentLoader:
    """Loads raw documents from any supported source.

    Supported sources:
    - PDF files (``*.pdf``) — text extracted page-by-page, concatenated.
    - Plain text files (``*.txt``)
    - Markdown files (``*.md``)
    - Word documents (``*.docx``)
    - HTTP/HTTPS URLs — HTML fetched and parsed via BeautifulSoup.
    - Directories — all supported files are loaded recursively.
    """

    def load(self, source: str) -> list[RawDocument]:
        """Load documents from *source* (file path, directory, or URL).

        Returns a list of :class:`RawDocument` objects.
        Raises :class:`ValueError` for unsupported file extensions.
        """
        if source.startswith("http://") or source.startswith("https://"):
            return self._load_url(source)

        path = Path(source)
        if path.is_dir():
            return self._load_directory(str(path))

        # Dispatch by extension first so we get a clear "unsupported extension" error
        # even if the file does not exist yet (e.g. in tests).
        ext = path.suffix.lower()
        if ext and ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file extension '{ext}'. "
                f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )

        if path.is_file():
            if ext == ".pdf":
                return self._load_pdf(str(path))
            elif ext == ".txt":
                return self._load_txt(str(path))
            elif ext == ".md":
                return self._load_md(str(path))
            elif ext == ".docx":
                return self._load_docx(str(path))

        raise ValueError(f"Source '{source}' is not a file, directory, or supported URL.")

    # ------------------------------------------------------------------
    # Private loader methods
    # ------------------------------------------------------------------

    def _load_pdf(self, path: str) -> list[RawDocument]:
        """Load a PDF file, concatenating all pages into one RawDocument.

        If page-level attribution is needed downstream, the ``page_number``
        field can be set on individual :class:`~src.models.Chunk` objects
        during chunking.
        """
        import pymupdf  # noqa: PLC0415 — imported lazily to avoid cost when not needed

        doc = pymupdf.open(path)
        pages_text: list[str] = []
        for page in doc:
            pages_text.append(page.get_text())
        content = "\n".join(pages_text)
        return [RawDocument(content=content, source=path, doc_type="pdf")]

    def _load_txt(self, path: str) -> list[RawDocument]:
        """Load a plain-text file."""
        content = Path(path).read_text(encoding="utf-8", errors="replace")
        return [RawDocument(content=content, source=path, doc_type="txt")]

    def _load_md(self, path: str) -> list[RawDocument]:
        """Load a Markdown file (treated as plain text)."""
        content = Path(path).read_text(encoding="utf-8", errors="replace")
        return [RawDocument(content=content, source=path, doc_type="md")]

    def _load_docx(self, path: str) -> list[RawDocument]:
        """Load a Microsoft Word .docx file."""
        import docx  # noqa: PLC0415 — imported lazily

        document = docx.Document(path)
        content = "\n".join(paragraph.text for paragraph in document.paragraphs)
        return [RawDocument(content=content, source=path, doc_type="docx")]

    def _load_url(self, url: str) -> list[RawDocument]:
        """Fetch an HTTP/HTTPS URL and extract visible text via BeautifulSoup."""
        import httpx  # noqa: PLC0415 — imported lazily
        from bs4 import BeautifulSoup  # noqa: PLC0415

        response = httpx.get(url, follow_redirects=True, timeout=30)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        content = soup.get_text(separator="\n", strip=True)
        return [RawDocument(content=content, source=url, doc_type="url")]

    def _load_directory(self, directory: str) -> list[RawDocument]:
        """Recursively load all supported files from *directory*."""
        results: list[RawDocument] = []
        for file_path in sorted(Path(directory).rglob("*")):
            if not file_path.is_file():
                continue
            ext = file_path.suffix.lower()
            if ext not in SUPPORTED_EXTENSIONS:
                logger.warning(
                    "Skipping unsupported file: %s (extension '%s' not supported)",
                    file_path,
                    ext,
                )
                continue
            try:
                docs = self.load(str(file_path))
                results.extend(docs)
            except Exception as exc:  # noqa: BLE001
                logger.warning("Failed to load '%s': %s", file_path, exc)
        return results
